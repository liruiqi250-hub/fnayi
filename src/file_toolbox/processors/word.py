from pathlib import Path
import logging
from typing import Callable

from docx import Document
from docx.oxml.ns import qn

from file_toolbox.translator import Translator


logger = logging.getLogger(__name__)
DRAWING_TAG = qn("w:drawing")


def _is_docx(path: Path) -> bool:
    return path.suffix.lower() == ".docx"


def _read_doc_text_via_ole(path: Path) -> str:
    """Fallback: extract raw text from .doc via olefile (basic, no formatting)."""
    import olefile

    ole = olefile.OleFileIO(path)
    try:
        if ole.exists("WordDocument"):
            data = ole.openstream("WordDocument").read()
            try:
                text = data.decode("utf-16-le", errors="ignore")
            except Exception:
                text = data.decode("latin-1", errors="ignore")
            text = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
            return text.strip()
        else:
            for stream_name in ole.listdir():
                data = ole.openstream(stream_name).read()
                try:
                    text = data.decode("utf-16-le", errors="ignore")
                    text = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
                    if len(text) > 50:
                        return text.strip()
                except Exception:
                    continue
            return ""
    finally:
        ole.close()


def _read_doc_text(path: Path) -> str:
    """Read text from a .doc file using pywin32 / Word COM automation."""
    word = None
    doc = None
    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(path.resolve()))
        text = doc.Content.Text
        return text
    except Exception:
        logger.warning(
            "win32com/Word COM failed for %s, falling back to olefile", path, exc_info=True
        )
        return _read_doc_text_via_ole(path)
    finally:
        try:
            if doc is not None:
                doc.Close(SaveChanges=0)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass


def _has_images(para) -> bool:
    """Check if a paragraph contains embedded images/drawings."""
    for run in para.runs:
        if not run.text.strip():
            if run._element.findall(".//" + DRAWING_TAG):
                return True
    return False


def _translate_paragraph(para, translator, target_language, source_language, progress=None):
    """Translate paragraph text while preserving run-level formatting
    (font, size, bold, color, etc.) and images."""
    if not para.text.strip():
        return

    full_text = para.text
    text_runs = [r for r in para.runs if r.text.strip()]
    if not text_runs:
        return

    translated = translator.translate(full_text, target_language, source_language)

    if _has_images(para):
        text_runs[0].text = translated
        for r in text_runs[1:]:
            r.text = ""
    else:
        # Preserve formatting by keeping the first run's style
        # and clearing the rest, instead of nuking all runs with para.text
        text_runs[0].text = translated
        for r in text_runs[1:]:
            r.text = ""


def _translate_cell(cell, translator, target_language, source_language, progress=None):
    """Translate all paragraphs in a table cell while preserving images."""
    for p in cell.paragraphs:
        _translate_paragraph(p, translator, target_language, source_language)


def translate_docx(source: Path, output_dir: Path, translator: Translator) -> Path:
    return translate_docx_to_language(source, output_dir, translator)


def translate_docx_to_language(
    source: Path,
    output_dir: Path,
    translator: Translator,
    target_language: str = "English",
    source_language: str = "Auto-detect",
    target_suffix: str = "EN",
    progress_callback: Callable[[int, int], None] | None = None,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    output = output_dir / f"{source.stem}_{target_suffix}.docx"

    if _is_docx(source):
        document = Document(source)
        items = []
        for p in document.paragraphs:
            if p.text.strip():
                items.append(("para", p))
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        items.append(("cell", cell))
        total = len(items)
        for idx, (kind, obj) in enumerate(items, 1):
            if kind == "para":
                _translate_paragraph(obj, translator, target_language, source_language)
            else:
                _translate_cell(obj, translator, target_language, source_language)
            if progress_callback:
                progress_callback(idx, total)
        document.save(output)
    else:
        raw_text = _read_doc_text(source)
        if not raw_text.strip():
            raise RuntimeError(f"无法从 {source.name} 提取文本，可能不是有效的 Word 文档。")
        document = Document()
        lines = [line.strip() for line in raw_text.split("\n") if line.strip()]
        total = len(lines)
        for idx, line in enumerate(lines, 1):
            document.add_paragraph(translator.translate(line, target_language, source_language))
            if progress_callback:
                progress_callback(idx, total)
        document.save(output)
    return output
