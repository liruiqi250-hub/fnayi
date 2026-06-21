from pathlib import Path

from docx import Document

from file_toolbox.processors.word import translate_docx
from conftest import FakeTranslator


def test_translate_docx_creates_english_copy(tmp_path: Path):
    source = tmp_path / "manual.docx"
    doc = Document()
    doc.add_heading("Product Manual", level=1)
    doc.add_paragraph("This is a sample product.")
    doc.save(source)

    output = translate_docx(source, tmp_path / "translated", FakeTranslator())

    assert output.name == "manual_EN.docx"
    translated = Document(output)
    text = "\n".join(p.text for p in translated.paragraphs)
    assert "[EN] Product Manual" in text
    assert "[EN] This is a sample product." in text
